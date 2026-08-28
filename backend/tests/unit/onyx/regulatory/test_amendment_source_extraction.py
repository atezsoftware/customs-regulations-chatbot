import ssl
import subprocess
import sys
import zipfile
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document

from onyx.regulatory.amendments import source_extraction
from onyx.regulatory.amendments.source_extraction import (
    AmendmentSourceExtractionError,
    extract_amendment_html,
    extract_amendment_pdf,
    fetch_and_extract_amendment_url,
)


class _Response:
    def __init__(self, content: bytes, content_type: str | None = None) -> None:
        self.headers = {"content-type": content_type or ""}
        self.url = "https://example.gov/update"
        self._content = content

    def iter_content(self, chunk_size: int) -> list[bytes]:
        return [
            self._content[index : index + chunk_size]
            for index in range(0, len(self._content), chunk_size)
        ]

    def raise_for_status(self) -> None:
        return None


def test_source_extraction_imports_without_playwright() -> None:
    script = """
import builtins
import sys

sys.path.insert(0, "backend")

original_import = builtins.__import__

def block_playwright(name, *args, **kwargs):
    if name == \"playwright\" or name.startswith(\"playwright.\"):
        raise ModuleNotFoundError(\"No module named 'playwright'\")
    return original_import(name, *args, **kwargs)

builtins.__import__ = block_playwright
import onyx.regulatory.amendments.source_extraction
"""

    result = subprocess.run(
        [sys.executable, "-c", script],
        capture_output=True,
        text=True,
        check=False,
    )

    assert result.returncode == 0, result.stderr


def test_extract_amendment_html_removes_page_chrome() -> None:
    result = extract_amendment_html(
        b"<html><nav>Menu</nav><main><h1>TEBLIG</h1><p>MADDE 1- Yeni metin.</p></main><script>bad()</script></html>",
        "text/html; charset=utf-8",
    )

    assert result == "TEBLIG\n\nMADDE 1- Yeni metin."


def test_extract_amendment_html_honors_declared_turkish_charset() -> None:
    content = (
        b'<html><head><meta http-equiv=Content-Type content="text/html; '
        b'charset=Windows-1254"></head><main>16 Kas\xfdm TEBL\xdd\xd0</main></html>'
    )

    result = extract_amendment_html(content, "text/html")

    assert result == "16 Kasım TEBLİĞ"


def test_extract_amendment_pdf_rejects_empty_text(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(
        source_extraction, "extract_file_text", lambda *_args, **_kwargs: ""
    )

    with pytest.raises(AmendmentSourceExtractionError, match="text-searchable"):
        extract_amendment_pdf(b"%PDF-1.7", "update.pdf")


def test_extract_amendment_docx_returns_normalized_text() -> None:
    document_bytes = BytesIO()
    document = Document()
    document.add_heading("ÜCRETLİ KARAYOLLARI", level=1)
    document.add_paragraph("MADDE 1- Geçiş ücretleri tahsil edilir.")
    document.save(document_bytes)

    result = source_extraction.extract_amendment_docx(
        document_bytes.getvalue(), "değişiklik.docx"
    )

    assert result == (
        "# ÜCRETLİ KARAYOLLARI\n\nMADDE 1- Geçiş ücretleri tahsil edilir."
    )


def test_extract_amendment_docx_rejects_non_docx_filename() -> None:
    document_bytes = BytesIO()
    document = Document()
    document.add_paragraph("MADDE 1- Yeni metin.")
    document.save(document_bytes)

    with pytest.raises(AmendmentSourceExtractionError, match="only .docx"):
        source_extraction.extract_amendment_docx(
            document_bytes.getvalue(), "değişiklik.doc"
        )


def test_extract_amendment_docx_rejects_oversized_file() -> None:
    with pytest.raises(AmendmentSourceExtractionError, match="too large"):
        source_extraction.extract_amendment_docx(
            b"x" * (source_extraction.MAX_AMENDMENT_SOURCE_BYTES + 1),
            "değişiklik.docx",
        )


def test_extract_amendment_docx_rejects_invalid_document() -> None:
    with pytest.raises(AmendmentSourceExtractionError, match="valid Word"):
        source_extraction.extract_amendment_docx(
            b"not a Word document", "değişiklik.docx"
        )


def test_extract_amendment_docx_rejects_corrupt_word_xml() -> None:
    document_bytes = BytesIO()
    with zipfile.ZipFile(document_bytes, "w") as archive:
        archive.writestr("[Content_Types].xml", b"not xml")
        archive.writestr("word/document.xml", b"not xml")

    with pytest.raises(AmendmentSourceExtractionError, match="could not be read"):
        source_extraction.extract_amendment_docx(
            document_bytes.getvalue(), "bozuk.docx"
        )


def test_extract_amendment_docx_rejects_excessive_expanded_size(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(
        source_extraction,
        "_MAX_AMENDMENT_DOCX_EXPANDED_BYTES",
        100,
        raising=False,
    )
    document_bytes = BytesIO()
    document = Document()
    document.add_paragraph("MADDE 1- Yeni metin.")
    document.save(document_bytes)

    with pytest.raises(AmendmentSourceExtractionError, match="expands beyond"):
        source_extraction.extract_amendment_docx(
            document_bytes.getvalue(), "değişiklik.docx"
        )


def test_extract_amendment_docx_rejects_too_many_archive_entries(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(source_extraction, "MAX_ARCHIVE_ENTRIES", 5, raising=False)
    document_bytes = BytesIO()
    document = Document()
    document.add_paragraph("MADDE 1- Yeni metin.")
    document.save(document_bytes)

    with pytest.raises(AmendmentSourceExtractionError, match="too many entries"):
        source_extraction.extract_amendment_docx(
            document_bytes.getvalue(), "değişiklik.docx"
        )


def test_extract_amendment_docx_rejects_oversized_xml_part(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(
        source_extraction, "_MAX_AMENDMENT_DOCX_XML_BYTES", 5, raising=False
    )
    document_bytes = BytesIO()
    with zipfile.ZipFile(document_bytes, "w") as archive:
        archive.writestr("[Content_Types].xml", b"<x/>")
        archive.writestr("word/document.xml", b"<doc/>")

    with pytest.raises(AmendmentSourceExtractionError, match="XML part"):
        source_extraction.extract_amendment_docx(
            document_bytes.getvalue(), "buyuk-xml.docx"
        )


def test_extract_amendment_docx_rejects_suspicious_compression_ratio() -> None:
    document_bytes = BytesIO()
    with zipfile.ZipFile(
        document_bytes, "w", compression=zipfile.ZIP_DEFLATED
    ) as archive:
        archive.writestr("[Content_Types].xml", b"not xml")
        archive.writestr("word/document.xml", b"not xml")
        archive.writestr("word/padding.bin", b"0" * (1024 * 1024))

    with pytest.raises(AmendmentSourceExtractionError, match="compression ratio"):
        source_extraction.extract_amendment_docx(
            document_bytes.getvalue(), "sikistirilmis.docx"
        )


def test_fetch_url_detects_pdf_by_signature(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(
        source_extraction,
        "ssrf_safe_get",
        lambda *_args, **_kwargs: _Response(b"%PDF-1.7"),
    )
    monkeypatch.setattr(
        source_extraction,
        "extract_amendment_pdf",
        lambda *_args, **_kwargs: "MADDE 1- Yeni metin.",
    )

    result = fetch_and_extract_amendment_url("https://example.gov/update")

    assert result.text == "MADDE 1- Yeni metin."
    assert result.source_type == "pdf"
    assert result.display_name == "update"


def test_fetch_resmi_gazete_url_uses_augmented_ca_bundle(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    request_kwargs: dict[str, object] = {}

    def fake_get(*_args: object, **kwargs: object) -> _Response:
        request_kwargs.update(kwargs)
        return _Response(b"<main>Resmi Gazete metni</main>", "text/html")

    monkeypatch.setattr(source_extraction, "ssrf_safe_get", fake_get)

    fetch_and_extract_amendment_url(
        "https://www.resmigazete.gov.tr/eskiler/2024/11/20241116-2.htm"
    )

    assert "verify" in request_kwargs
    ca_bundle_path = Path(str(request_kwargs["verify"]))
    assert ca_bundle_path.is_file()
    ssl.create_default_context(cafile=ca_bundle_path)


def test_fetch_url_reports_tls_verification_failure(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    def fail_tls(*_args: object, **_kwargs: object) -> _Response:
        raise source_extraction.requests.exceptions.SSLError(
            "certificate verify failed"
        )

    monkeypatch.setattr(source_extraction, "ssrf_safe_get", fail_tls)

    with pytest.raises(
        AmendmentSourceExtractionError,
        match="TLS certificate chain could not be verified",
    ):
        fetch_and_extract_amendment_url("https://example.gov/update.htm")


def test_fetch_url_rejects_oversized_stream(monkeypatch: pytest.MonkeyPatch) -> None:
    monkeypatch.setattr(
        source_extraction,
        "ssrf_safe_get",
        lambda *_args, **_kwargs: _Response(
            b"x" * (source_extraction.MAX_AMENDMENT_SOURCE_BYTES + 1),
            "text/html",
        ),
    )

    with pytest.raises(AmendmentSourceExtractionError, match="too large"):
        fetch_and_extract_amendment_url("https://example.gov/update.htm")
