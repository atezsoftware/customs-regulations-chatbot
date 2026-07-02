"""
Standalone regulatory document chunking.

This module is intentionally not wired into the main indexing pipeline yet. It
builds inspectable chunks with locator metadata so we can evaluate boundaries on
real DOCX files before changing retrieval behavior.
"""

from __future__ import annotations

import hashlib
import re
import unicodedata
from dataclasses import asdict, dataclass, field, replace
from pathlib import Path
from typing import Any

from fs_explorer_shared.fs import SUPPORTED_EXTENSIONS

_PARSE_ERROR_PREFIXES: tuple[str, ...] = (
    "Error parsing ",
    "Unsupported file extension",
    "No such file:",
)

_DATE_DMY_RE = re.compile(r"(?<!\d)(\d{1,2})[./](\d{1,2})[./](\d{4})(?!\d)")
_DATE_YMD_RE = re.compile(r"(?<!\d)(\d{4})[-/.](\d{1,2})[-/.](\d{1,2})(?!\d)")
_HEADER_NUMBER_RE = re.compile(r"(?i)\bSayı\s*[:：]\s*([^\n]+)")
_NORM_NUMBER_RE = re.compile(r"\b(\d{4})\s*/\s*(\d{1,6})\b")
_FILENAME_NUMBER_RE = re.compile(
    r"(?i)\b(?:genelge|bkk|ck|kugm|gkgm|rykgm|dhdb)[_-](\d{4})[-_/](\d{1,6})\b"
)

_ARTICLE_RE = re.compile(
    r"^(?P<label>(?:gecici|geçici)\s+madde|(?:mukerrer|mükerrer)\s+madde|madde)"
    r"\s+(?P<num>\d+[a-z]?)\b\s*(?:[-–—:]\s*)?(?P<rest>.*)$",
    flags=re.IGNORECASE,
)
_SECTION_ORDINAL_PATTERN = (
    r"(?:birinci|ikinci|ucuncu|üçüncü|dorduncu|dördüncü|besinci|beşinci|"
    r"altinci|altıncı|yedinci|sekizinci|dokuzuncu|onuncu|on\s*birinci|"
    r"on\s*ikinci|[ivxlcdm]+|\d+)"
)
_SECTION_RE = re.compile(
    rf"^(?:(?:{_SECTION_ORDINAL_PATTERN})\s+(?:bolum|bölüm|kisim|kısım)\b|"
    rf"(?:bolum|bölüm|kisim|kısım)\s+(?:{_SECTION_ORDINAL_PATTERN})\b)",
    flags=re.IGNORECASE,
)
_APPENDIX_RE = re.compile(
    r"^(?P<label>ek)\s*[-: ]\s*(?P<num>[a-z0-9ivxlcdm]+)\b|^(?P<label2>ek)\b$",
    flags=re.IGNORECASE,
)
_NUMBERED_SECTION_RE = re.compile(r"^(?P<num>\d{1,3})[.)]\s+(?P<title>.+?)(?::)?$")
_PARAGRAPH_RE = re.compile(r"^\(?(?P<num>\d{1,3})[.)](?!\d)\s*")
_REPEATED_PARAGRAPH_RE = re.compile(
    r"^(?P<marker>(?:mükerrer|mukerrer)\s*[-–—]?\s*(?P<num>\d{1,3}))"
    r"\s*[,.:;-]?\s*",
    flags=re.IGNORECASE,
)
_CLAUSE_RE = re.compile(r"^\(?(?P<label>[a-zçğıöşü])\)\s+", flags=re.IGNORECASE)
_SUBCLAUSE_RE = re.compile(r"^\((?P<label>[ivxlcdm]+|\d+)\)\s+", flags=re.IGNORECASE)
_MARKDOWN_TABLE_RE = re.compile(r"^\s*\|.*\|\s*$")
_ROMAN_LABELS = {
    "i",
    "ii",
    "iii",
    "iv",
    "v",
    "vi",
    "vii",
    "viii",
    "ix",
    "x",
    "xi",
    "xii",
    "xiii",
    "xiv",
    "xv",
    "xvi",
    "xvii",
    "xviii",
    "xix",
    "xx",
}

_DOCUMENT_TYPE_ORDER: tuple[tuple[str, tuple[str, ...]], ...] = (
    ("kanun", ("kanun",)),
    ("sozlesme", ("sozlesme", "sözlesme", "sözleşme", "andlasma", "andlaşma")),
    ("yonetmelik", ("yonetmelik", "yönetmelik")),
    ("teblig", ("teblig", "tebliğ")),
    ("genelge", ("genelge",)),
    ("yonerge", ("yonerge", "yönerge")),
    ("karar", ("karar", "bkk", "cumhurbaskani karari", "cumhurbaşkanı kararı")),
    ("protokol", ("protokol",)),
)


@dataclass(frozen=True)
class SourceBlock:
    """A normalized source block used by the standalone chunker."""

    block_id: int
    kind: str
    text: str
    start_char: int
    end_char: int


@dataclass(frozen=True)
class StructureNode:
    """A structural parent node used to locate chunks without becoming a chunk."""

    node_id: str
    parent_id: str | None
    node_type: str
    label: str
    level: int
    order: int
    source_block: int | None = None


@dataclass(frozen=True)
class DocumentMetadata:
    """Small, generic document metadata used by chunk locator metadata."""

    doc_id: str
    source_file: str
    source_path: str
    file_date: str | None
    document_date: str | None
    document_number: str | None
    document_type: str
    title: str | None
    parser: str


@dataclass(frozen=True)
class ChunkMetadata:
    """Locator metadata for a chunk.

    Some fields here exist only so `IndexingPipeline` can derive the stable
    top-level `core_chunks` columns (`id`, `position`, `start_char`,
    `end_char`) — they are intentionally left out of `to_storage_dict()`
    since persisting them again inside the `metadata` JSON column would just
    duplicate those columns (and `chunk_id`/`doc_id` here are this chunker's
    own internal hashes, distinct from and easily confused with the real
    `core_chunks.id`/`core_documents.id` the pipeline actually assigns).
    """

    chunk_id: str
    doc_id: str
    source_file: str
    source_path: str
    document_date: str | None
    document_number: str | None
    document_type: str
    chunk_type: str
    chunk_order: int
    parent_id: str | None
    parent_ids: list[str]
    parent_path: list[str]
    heading_path: list[str]
    article_no: str | None = None
    article_title: str | None = None
    paragraph_no: str | None = None
    clause_label: str | None = None
    subclause_label: str | None = None
    appendix_label: str | None = None
    table_index: int | None = None
    table_row_index: int | None = None
    source_start_char: int = 0
    source_end_char: int = 0
    source_start_block: int = 0
    source_end_block: int = 0

    def to_storage_dict(self) -> dict[str, Any]:
        """Trimmed projection persisted to `core_chunks.metadata`.

        Excludes fields that only exist for internal id derivation
        (`chunk_id`, `doc_id`, `chunk_type`, `chunk_order`, `source_*char`,
        `source_*block`) or that duplicate `heading_path`
        (`parent_id`/`parent_ids`/`parent_path`) or the document's own
        `relative_path` (`source_file`/`source_path`) — all already
        available as top-level `core_chunks`/`core_documents` columns.
        """
        return {
            "document_date": self.document_date,
            "document_number": self.document_number,
            "document_type": self.document_type,
            "heading_path": self.heading_path,
            "article_no": self.article_no,
            "article_title": self.article_title,
            "paragraph_no": self.paragraph_no,
            "clause_label": self.clause_label,
            "subclause_label": self.subclause_label,
            "appendix_label": self.appendix_label,
            "table_index": self.table_index,
            "table_row_index": self.table_row_index,
        }


@dataclass(frozen=True)
class RegulatoryChunk:
    """Chunk text plus its locator metadata."""

    text: str
    metadata: ChunkMetadata

    def to_dict(self) -> dict[str, Any]:
        payload = asdict(self.metadata)
        payload["text"] = self.text
        return payload


@dataclass(frozen=True)
class ChunkedDocument:
    """Full standalone chunking output for frontend inspection."""

    metadata: DocumentMetadata
    content: str
    structure: list[StructureNode]
    blocks: list[SourceBlock]
    chunks: list[RegulatoryChunk]
    warnings: list[str]

    def to_dict(self) -> dict[str, Any]:
        return {
            "metadata": asdict(self.metadata),
            "content": self.content,
            "structure": [asdict(node) for node in self.structure],
            "blocks": [asdict(block) for block in self.blocks],
            "chunks": [chunk.to_dict() for chunk in self.chunks],
            "warnings": list(self.warnings),
            "stats": {
                "block_count": len(self.blocks),
                "chunk_count": len(self.chunks),
                "content_chars": len(self.content),
                "chunk_chars": sum(len(chunk.text) for chunk in self.chunks),
            },
        }


@dataclass
class _ChunkDraft:
    chunk_type: str
    blocks: list[SourceBlock]
    parent_id: str | None
    parent_ids: list[str]
    parent_path: list[str]
    heading_path: list[str]
    article_no: str | None = None
    article_title: str | None = None
    paragraph_no: str | None = None
    clause_label: str | None = None
    subclause_label: str | None = None
    appendix_label: str | None = None
    table_index: int | None = None
    table_row_index: int | None = None
    warnings: list[str] = field(default_factory=list)


class RegulatoryChunker:
    """Generic Turkish regulatory-text chunker."""

    def __init__(
        self, *, max_chunk_chars: int = 2400, min_chunk_chars: int = 120
    ) -> None:
        if max_chunk_chars < 500:
            raise ValueError("max_chunk_chars must be >= 500")
        if min_chunk_chars < 0:
            raise ValueError("min_chunk_chars must be >= 0")
        self.max_chunk_chars = max_chunk_chars
        self.min_chunk_chars = min_chunk_chars

    def chunk_file(
        self, file_path: str, *, root_path: str | None = None
    ) -> ChunkedDocument:
        path = Path(file_path).expanduser().resolve()
        if not path.exists() or not path.is_file():
            raise ValueError(f"No such file: {path}")
        if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file extension: {path.suffix}. "
                f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
            )

        content, parser, warnings = _parse_with_fallback(path)
        blocks = _markdown_to_blocks(content)
        if not blocks:
            warnings.append("No inspectable text blocks were produced.")

        metadata = _infer_document_metadata(
            path=path,
            root_path=Path(root_path).expanduser().resolve() if root_path else None,
            content=content,
            parser=parser,
        )
        chunks, structure = self.chunk_blocks_with_structure(blocks, metadata=metadata)
        if not chunks and blocks:
            warnings.append("No chunks were produced from non-empty blocks.")

        return ChunkedDocument(
            metadata=metadata,
            content=content,
            structure=structure,
            blocks=blocks,
            chunks=chunks,
            warnings=warnings,
        )

    def chunk_text(
        self,
        text: str,
        *,
        source_path: str = "<memory>",
        source_file: str = "<memory>",
        parser: str = "memory",
    ) -> ChunkedDocument:
        content = text.strip()
        metadata = _infer_document_metadata(
            path=Path(source_file),
            root_path=None,
            content=content,
            parser=parser,
            source_path_override=source_path,
        )
        blocks = _markdown_to_blocks(content)
        chunks, structure = self.chunk_blocks_with_structure(blocks, metadata=metadata)
        return ChunkedDocument(
            metadata=metadata,
            content=content,
            structure=structure,
            blocks=blocks,
            chunks=chunks,
            warnings=[],
        )

    def chunk_blocks(
        self,
        blocks: list[SourceBlock],
        *,
        metadata: DocumentMetadata,
    ) -> list[RegulatoryChunk]:
        chunks, _ = self.chunk_blocks_with_structure(blocks, metadata=metadata)
        return chunks

    def chunk_blocks_with_structure(
        self,
        blocks: list[SourceBlock],
        *,
        metadata: DocumentMetadata,
    ) -> tuple[list[RegulatoryChunk], list[StructureNode]]:
        drafts, structure = self._build_drafts(blocks, metadata=metadata)
        drafts = _drop_non_atomic_parent_drafts(drafts, metadata)
        chunks: list[RegulatoryChunk] = []
        for draft in drafts:
            for split in self._split_oversized_draft(draft):
                if _is_non_chunk_draft(split, metadata):
                    continue
                if _draft_text(split).strip():
                    chunks.append(self._finalize_chunk(split, metadata, len(chunks)))

        # The title-vs-content heuristic in `_is_non_chunk_draft` can filter out
        # every draft when the whole document is just one short, title-like
        # line — better to keep that line as a chunk than to index nothing.
        if not chunks and blocks:
            non_empty_drafts = [d for d in drafts if _draft_text(d).strip()]
            if non_empty_drafts:
                chunks.append(self._finalize_chunk(non_empty_drafts[0], metadata, 0))

        return chunks, structure

    def _build_drafts(
        self, blocks: list[SourceBlock], *, metadata: DocumentMetadata
    ) -> tuple[list[_ChunkDraft], list[StructureNode]]:
        drafts: list[_ChunkDraft] = []
        structure: list[StructureNode] = []
        current: _ChunkDraft | None = None
        pending_title: str | None = None
        pending_title_block: SourceBlock | None = None
        current_article_no: str | None = None
        current_article_title: str | None = None
        current_appendix: str | None = None
        table_index = 0

        document_label = metadata.title or metadata.source_file
        document_node = _make_structure_node(
            metadata=metadata,
            node_type="document",
            label=document_label,
            parent_id=None,
            level=0,
            order=len(structure),
            source_block=None,
        )
        structure.append(document_node)
        stack: list[StructureNode] = [document_node]

        def promote_parent_context(draft: _ChunkDraft) -> bool:
            context = _preamble_context_label(draft)
            if context is None:
                return False
            label, source_block = context
            push_node(
                "preamble",
                label,
                level=1,
                source_block=source_block,
            )
            return True

        def flush() -> None:
            nonlocal current
            if current is not None and current.blocks:
                if not promote_parent_context(current):
                    drafts.append(current)
            current = None

        def push_node(
            node_type: str,
            label: str,
            *,
            level: int,
            source_block: int | None,
        ) -> StructureNode:
            nonlocal stack
            clean_label = _clean_inline_markdown(label)
            if not clean_label:
                return stack[-1]
            if _same_label(clean_label, stack[-1].label):
                return stack[-1]
            while len(stack) > 1 and stack[-1].level >= level:
                stack.pop()
            parent = stack[-1] if stack else None
            node = _make_structure_node(
                metadata=metadata,
                node_type=node_type,
                label=clean_label,
                parent_id=parent.node_id if parent else None,
                level=level,
                order=len(structure),
                source_block=source_block,
            )
            structure.append(node)
            stack.append(node)
            return node

        def parent_values() -> tuple[str | None, list[str], list[str]]:
            parent_ids = [node.node_id for node in stack]
            parent_path = [node.label for node in stack]
            return stack[-1].node_id if stack else None, parent_ids, parent_path

        def apply_pending_heading(*, level: int = 3) -> None:
            nonlocal pending_title, pending_title_block
            if pending_title is None:
                return
            if not _same_label(pending_title, document_label):
                push_node(
                    "heading",
                    pending_title,
                    level=level,
                    source_block=(
                        pending_title_block.block_id if pending_title_block else None
                    ),
                )
            pending_title = None
            pending_title_block = None

        def apply_article_context_from_current() -> None:
            if current is None or current.chunk_type != "article":
                return
            context = _article_context_label(current)
            if context is None:
                return
            label, source_block = context
            push_node(
                "context",
                label,
                level=_article_context_level(stack),
                source_block=source_block,
            )

        for block in blocks:
            classification = _classify_block(block)
            clean_text = _clean_inline_markdown(block.text)

            if classification["kind"] == "section":
                flush()
                section_level = _section_level(clean_text)
                if pending_title is not None:
                    apply_pending_heading(level=max(1, section_level - 1))
                push_node(
                    "section",
                    clean_text,
                    level=section_level,
                    source_block=block.block_id,
                )
                current_article_no = None
                current_article_title = None
                continue

            if classification["kind"] == "appendix":
                flush()
                current_appendix = str(classification.get("label") or clean_text)
                apply_pending_heading(level=1)
                push_node(
                    "appendix",
                    current_appendix,
                    level=1,
                    source_block=block.block_id,
                )
                current_article_no = None
                current_article_title = None
                continue

            if classification["kind"] == "article":
                article_title_candidate = pending_title
                flush()
                current_article_no = str(classification.get("article_no") or "")
                current_article_title = _article_title_from_rest(
                    str(classification.get("rest") or ""), article_title_candidate
                )
                article_heading = str(
                    classification.get("article_heading")
                    or f"MADDE {current_article_no}"
                )
                if current_article_title:
                    article_heading = f"{article_heading} - {current_article_title}"
                if pending_title and current_article_title != pending_title:
                    pending_level = (
                        stack[-1].level + 1
                        if stack
                        and stack[-1].node_type == "heading"
                        and stack[-1].level >= 3
                        else 3
                    )
                    apply_pending_heading(level=pending_level)
                else:
                    pending_title = None
                    pending_title_block = None
                article_level = _child_level(stack, 4)
                push_node(
                    "article",
                    article_heading,
                    level=article_level,
                    source_block=block.block_id,
                )
                inline_unit = _article_unit(
                    str(classification.get("rest") or ""),
                    block.text,
                )
                if inline_unit is not None and not _is_structural_article_unit_heading(
                    inline_unit
                ):
                    push_node(
                        str(inline_unit["kind"]),
                        _article_unit_parent_label(
                            inline_unit, str(classification.get("rest") or "")
                        ),
                        level=article_level + 1,
                        source_block=block.block_id,
                    )
                    parent_id, parent_ids, parent_path = parent_values()
                    unit_kind = str(inline_unit["kind"])
                    unit_label = str(inline_unit["label"])
                    current = _ChunkDraft(
                        chunk_type=unit_kind,
                        blocks=[block],
                        parent_id=parent_id,
                        parent_ids=parent_ids,
                        parent_path=parent_path,
                        heading_path=parent_path,
                        article_no=current_article_no,
                        article_title=current_article_title,
                        paragraph_no=unit_label if unit_kind == "paragraph" else None,
                        clause_label=unit_label if unit_kind == "clause" else None,
                        appendix_label=current_appendix,
                    )
                    continue

                parent_id, parent_ids, parent_path = parent_values()
                current = _ChunkDraft(
                    chunk_type="article",
                    blocks=[block],
                    parent_id=parent_id,
                    parent_ids=parent_ids,
                    parent_path=parent_path,
                    heading_path=parent_path,
                    article_no=current_article_no,
                    article_title=current_article_title,
                    appendix_label=current_appendix,
                )
                continue

            # Not gated on `current_article_no` — a leading `N)`/`N.` paragraph
            # marker or `a)` clause marker means the same thing whether it's
            # inside a MADDE article or a non-MADDE numbered section/preamble
            # (protocols, conventions). `_article_unit_level()` and
            # `apply_article_context_from_current()` already degrade cleanly
            # with no article on the stack. Restricted to plain "paragraph"
            # classification so a heading (e.g. "**(a) TARİFLER**") or a
            # numbered-section title (e.g. "**1. ...:**") is never reinterpreted
            # as a clause/paragraph marker — those already have their own,
            # more specific handling below.
            article_unit = (
                _article_unit(clean_text, block.text)
                if classification["kind"] == "paragraph"
                else None
            )
            if article_unit is not None:
                if _is_structural_article_unit_heading(article_unit):
                    flush()
                    pending_title = clean_text
                    pending_title_block = block
                    continue

                apply_article_context_from_current()
                flush()
                apply_pending_heading(level=_article_context_level(stack))
                unit_level = _article_unit_level(stack, kind=str(article_unit["kind"]))
                push_node(
                    str(article_unit["kind"]),
                    _article_unit_parent_label(article_unit, clean_text),
                    level=unit_level,
                    source_block=block.block_id,
                )
                parent_id, parent_ids, parent_path = parent_values()
                unit_kind = str(article_unit["kind"])
                unit_label = str(article_unit["label"])
                current = _ChunkDraft(
                    chunk_type=unit_kind,
                    blocks=[block],
                    parent_id=parent_id,
                    parent_ids=parent_ids,
                    parent_path=parent_path,
                    heading_path=parent_path,
                    article_no=current_article_no,
                    article_title=current_article_title,
                    paragraph_no=unit_label if unit_kind == "paragraph" else None,
                    clause_label=unit_label if unit_kind == "clause" else None,
                    appendix_label=current_appendix,
                )
                continue

            if classification["kind"] == "numbered_section":
                flush()
                title = clean_text
                apply_pending_heading(level=2)
                push_node(
                    "numbered_section",
                    title,
                    level=3,
                    source_block=block.block_id,
                )
                parent_id, parent_ids, parent_path = parent_values()
                current = _ChunkDraft(
                    chunk_type="numbered_section",
                    blocks=[block],
                    parent_id=parent_id,
                    parent_ids=parent_ids,
                    parent_path=parent_path,
                    heading_path=parent_path,
                    appendix_label=current_appendix,
                )
                pending_title = None
                pending_title_block = None
                continue

            if block.kind == "table":
                flush()
                apply_pending_heading(level=3)
                table_index += 1
                parent_id, parent_ids, parent_path = parent_values()
                current = _ChunkDraft(
                    chunk_type="table",
                    blocks=[block],
                    parent_id=parent_id,
                    parent_ids=parent_ids,
                    parent_path=parent_path,
                    heading_path=parent_path,
                    article_no=current_article_no,
                    article_title=current_article_title,
                    appendix_label=current_appendix,
                    table_index=table_index,
                )
                flush()
                continue

            if classification["kind"] == "heading":
                incoming_unit = _article_unit(clean_text, block.text)
                if (
                    pending_title is not None
                    and incoming_unit is not None
                    and _is_structural_article_unit_heading(incoming_unit)
                ):
                    apply_pending_heading(level=3)
                if (
                    current is not None
                    and current.chunk_type == "article"
                    and not current.article_title
                    and len(current.blocks) == 1
                    and stack
                    and stack[-1].node_type == "article"
                ):
                    # "MADDE N" and its title sit on separate lines (no inline
                    # rest). Without this, flush() below would close MADDE N
                    # as a title-only chunk and leave its structure node on
                    # the stack forever (never popped, since nothing pushes a
                    # replacement at its level) — every later article would
                    # then nest one level deeper under it instead of becoming
                    # its sibling.
                    current.article_title = clean_text
                    current.blocks.append(block)
                    updated_node = replace(
                        stack[-1], label=f"{stack[-1].label} - {clean_text}"
                    )
                    stack[-1] = updated_node
                    structure[-1] = updated_node
                    continue
                flush()
                pending_title = clean_text
                pending_title_block = block
                continue

            apply_pending_heading(
                level=_article_context_level(stack)
                if current_article_no is not None
                else 3
            )
            if current is None:
                parent_id, parent_ids, parent_path = parent_values()
                current = _ChunkDraft(
                    chunk_type="free_text",
                    blocks=[],
                    parent_id=parent_id,
                    parent_ids=parent_ids,
                    parent_path=parent_path,
                    heading_path=parent_path,
                    article_no=current_article_no,
                    article_title=current_article_title,
                    appendix_label=current_appendix,
                )
            assert current is not None
            draft = current

            if draft.chunk_type == "article":
                paragraph_no = _paragraph_no(clean_text)
                clause_label = _clause_label(clean_text)
                subclause_label = _subclause_label(clean_text)
                if paragraph_no and draft.paragraph_no is None:
                    draft.paragraph_no = paragraph_no
                if clause_label and draft.clause_label is None:
                    draft.clause_label = clause_label
                if subclause_label and draft.subclause_label is None:
                    draft.subclause_label = subclause_label

            draft.blocks.append(block)

        flush()
        return _merge_tiny_drafts(drafts, self.min_chunk_chars), structure

    def _split_oversized_draft(self, draft: _ChunkDraft) -> list[_ChunkDraft]:
        if len(_draft_text(draft)) <= self.max_chunk_chars:
            return [draft]

        if len(draft.blocks) <= 1:
            return [draft]

        splits: list[_ChunkDraft] = []
        current_blocks: list[SourceBlock] = []
        part = 1
        for block in draft.blocks:
            tentative = [*current_blocks, block]
            if current_blocks and len(_blocks_text(tentative)) > self.max_chunk_chars:
                splits.append(self._copy_draft_part(draft, current_blocks, part))
                part += 1
                current_blocks = [block]
            else:
                current_blocks = tentative

        if current_blocks:
            splits.append(self._copy_draft_part(draft, current_blocks, part))
        return splits

    @staticmethod
    def _copy_draft_part(
        draft: _ChunkDraft, blocks: list[SourceBlock], part: int
    ) -> _ChunkDraft:
        warnings = list(draft.warnings)
        warnings.append(f"Split from oversized {draft.chunk_type} chunk, part {part}.")
        # A size-driven cut isn't a structural unit in the source document —
        # don't fabricate a "Part N" heading for it. `chunk_order` plus
        # `source_start_char`/`source_end_char` already make each part
        # locatable, and the warning above makes the split traceable.
        heading_path = list(draft.heading_path)
        return _ChunkDraft(
            chunk_type=draft.chunk_type,
            blocks=list(blocks),
            parent_id=draft.parent_id,
            parent_ids=list(draft.parent_ids),
            parent_path=list(draft.parent_path),
            heading_path=heading_path,
            article_no=draft.article_no,
            article_title=draft.article_title,
            paragraph_no=(
                draft.paragraph_no
                if draft.chunk_type == "paragraph"
                else _paragraph_no(_clean_inline_markdown(blocks[0].text))
                or draft.paragraph_no
            ),
            clause_label=(
                draft.clause_label
                if draft.chunk_type == "clause"
                else _clause_label(_clean_inline_markdown(blocks[0].text))
                or draft.clause_label
            ),
            subclause_label=_subclause_label(_clean_inline_markdown(blocks[0].text))
            or draft.subclause_label,
            appendix_label=draft.appendix_label,
            table_index=draft.table_index,
            table_row_index=draft.table_row_index,
            warnings=warnings,
        )

    @staticmethod
    def _finalize_chunk(
        draft: _ChunkDraft,
        metadata: DocumentMetadata,
        order: int,
    ) -> RegulatoryChunk:
        text = _draft_text(draft).strip()
        start_block = draft.blocks[0]
        end_block = draft.blocks[-1]
        chunk_id = _stable_id(
            "chunk",
            (
                f"{metadata.source_path}:{order}:"
                f"{start_block.start_char}:{end_block.end_char}:{draft.chunk_type}"
            ),
        )
        chunk_metadata = ChunkMetadata(
            chunk_id=chunk_id,
            doc_id=metadata.doc_id,
            source_file=metadata.source_file,
            source_path=metadata.source_path,
            document_date=metadata.document_date,
            document_number=metadata.document_number,
            document_type=metadata.document_type,
            chunk_type=draft.chunk_type,
            chunk_order=order,
            parent_id=draft.parent_id,
            parent_ids=list(draft.parent_ids),
            parent_path=list(draft.parent_path),
            heading_path=list(draft.heading_path),
            article_no=draft.article_no,
            article_title=draft.article_title,
            paragraph_no=draft.paragraph_no,
            clause_label=draft.clause_label,
            subclause_label=draft.subclause_label,
            appendix_label=draft.appendix_label,
            table_index=draft.table_index,
            table_row_index=draft.table_row_index,
            source_start_char=start_block.start_char,
            source_end_char=end_block.end_char,
            source_start_block=start_block.block_id,
            source_end_block=end_block.block_id,
        )
        return RegulatoryChunk(text=text, metadata=chunk_metadata)


def _parse_with_fallback(path: Path) -> tuple[str, str, list[str]]:
    warnings: list[str] = []
    from ..document_parser import parse_file

    content = parse_file(str(path))
    if not content.startswith(_PARSE_ERROR_PREFIXES):
        return content.strip(), "docling_markdown", warnings

    warnings.append(content)
    if path.suffix.lower() != ".docx":
        return "", "parse_failed", warnings

    try:
        content = _parse_docx_with_python_docx(path)
        warnings.append("Used python-docx fallback parser.")
        return content.strip(), "python_docx_fallback", warnings
    except Exception as exc:
        warnings.append(f"python-docx fallback failed: {exc}")
        return "", "parse_failed", warnings


def _parse_docx_with_python_docx(path: Path) -> str:
    from docx import Document

    document = Document(str(path))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table_index, table in enumerate(document.tables, 1):
        parts.append(f"Table {table_index}")
        for row in table.rows:
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            if any(cells):
                parts.append("| " + " | ".join(cells) + " |")
    return "\n\n".join(parts)


def _markdown_to_blocks(content: str) -> list[SourceBlock]:
    blocks: list[SourceBlock] = []
    text = content.strip()
    if not text:
        return blocks

    lines = text.splitlines(keepends=True)
    offset = 0
    pending_table: list[tuple[str, int, int]] = []

    def flush_table() -> None:
        nonlocal pending_table
        if not pending_table:
            return
        table_text = "".join(item[0] for item in pending_table).strip()
        blocks.append(
            SourceBlock(
                block_id=len(blocks),
                kind="table",
                text=table_text,
                start_char=pending_table[0][1],
                end_char=pending_table[-1][2],
            )
        )
        pending_table = []

    for line in lines:
        start = offset
        end = offset + len(line)
        offset = end
        stripped = line.strip()
        if not stripped:
            flush_table()
            continue

        if _MARKDOWN_TABLE_RE.match(stripped):
            pending_table.append((line, start, end))
            continue

        flush_table()
        kind = "heading" if _looks_like_markdown_heading(stripped) else "paragraph"
        blocks.append(
            SourceBlock(
                block_id=len(blocks),
                kind=kind,
                text=stripped,
                start_char=start,
                end_char=end,
            )
        )

    flush_table()
    return blocks


def _infer_document_metadata(
    *,
    path: Path,
    root_path: Path | None,
    content: str,
    parser: str,
    source_path_override: str | None = None,
) -> DocumentMetadata:
    resolved = path.expanduser().resolve() if str(path) != "<memory>" else path
    if source_path_override is not None:
        source_path = source_path_override
    elif root_path is not None:
        try:
            source_path = str(resolved.relative_to(root_path))
        except ValueError:
            source_path = str(resolved)
    else:
        source_path = str(resolved)

    source_file = path.name
    file_date = _first_iso_date(path.name)
    header = "\n".join(content.splitlines()[:60])
    document_date = _first_iso_date(header) or file_date
    document_type = _infer_document_type(path.name, header)
    document_number = _infer_document_number(path.name, header)
    title = _infer_title(content)

    return DocumentMetadata(
        doc_id=_stable_id("doc", source_path),
        source_file=source_file,
        source_path=source_path,
        file_date=file_date,
        document_date=document_date,
        document_number=document_number,
        document_type=document_type,
        title=title,
        parser=parser,
    )


def _infer_document_number(filename: str, header: str) -> str | None:
    norm_match = _NORM_NUMBER_RE.search(header)
    if norm_match:
        return f"{norm_match.group(1)}/{norm_match.group(2)}"

    filename_match = _FILENAME_NUMBER_RE.search(filename)
    if filename_match:
        return f"{filename_match.group(1)}-{filename_match.group(2)}"

    header_match = _HEADER_NUMBER_RE.search(header)
    if header_match:
        value = _clean_inline_markdown(header_match.group(1))
        return value[:120] if value else None
    return None


def _infer_document_type(filename: str, header: str) -> str:
    folded_filename = _fold_text(filename)
    for doc_type, markers in _DOCUMENT_TYPE_ORDER:
        if any(_fold_text(marker) in folded_filename for marker in markers):
            return doc_type

    folded = _fold_text(header)
    for doc_type, markers in _DOCUMENT_TYPE_ORDER:
        if any(_fold_text(marker) in folded for marker in markers):
            return doc_type
    return "unknown"


def _infer_title(content: str) -> str | None:
    for line in content.splitlines():
        cleaned = _clean_inline_markdown(line)
        if not cleaned:
            continue
        folded = _fold_text(cleaned)
        if folded in {"t.c.", "tc"}:
            continue
        if folded.startswith("sayi") or folded.startswith("konu"):
            continue
        return cleaned[:240]
    return None


def _first_iso_date(text: str) -> str | None:
    dmy = _DATE_DMY_RE.search(text)
    if dmy:
        day, month, year = (int(dmy.group(1)), int(dmy.group(2)), int(dmy.group(3)))
        if _valid_date_parts(year, month, day):
            return f"{year:04d}-{month:02d}-{day:02d}"

    ymd = _DATE_YMD_RE.search(text)
    if ymd:
        year, month, day = (int(ymd.group(1)), int(ymd.group(2)), int(ymd.group(3)))
        if _valid_date_parts(year, month, day):
            return f"{year:04d}-{month:02d}-{day:02d}"
    return None


def _valid_date_parts(year: int, month: int, day: int) -> bool:
    return 1800 <= year <= 2200 and 1 <= month <= 12 and 1 <= day <= 31


def _make_structure_node(
    *,
    metadata: DocumentMetadata,
    node_type: str,
    label: str,
    parent_id: str | None,
    level: int,
    order: int,
    source_block: int | None,
) -> StructureNode:
    node_id = _stable_id(
        "node",
        f"{metadata.doc_id}:{parent_id or '<root>'}:{node_type}:{label}:{source_block}",
    )
    return StructureNode(
        node_id=node_id,
        parent_id=parent_id,
        node_type=node_type,
        label=label,
        level=level,
        order=order,
        source_block=source_block,
    )


def _same_label(left: str, right: str) -> bool:
    return _fold_text(_clean_inline_markdown(left)) == _fold_text(
        _clean_inline_markdown(right)
    )


def _section_level(text: str) -> int:
    folded = _fold_text(text)
    if "kisim" in folded:
        return 1
    if "bolum" in folded:
        return 2
    return 2


def _child_level(stack: list[StructureNode], base_level: int) -> int:
    """Level for a new article, anchored to the nearest section/appendix/document
    ancestor rather than the deepest stack entry.

    Articles are siblings of each other: MADDE 2 must sit at the same level as
    MADDE 1, not one level deeper. Anchoring on `stack[-1]` (the previous
    article's leftover context/heading nodes) made the level climb by one on
    every article, so `push_node`'s pop-while-level>=new-level never popped the
    prior article off the stack and `heading_path`/`parent_path` accumulated
    every previous MADDE instead of resetting per article.
    """
    anchor_types = {
        "document",
        "section",
        "appendix",
        "heading",
        "numbered_section",
        "preamble",
    }
    for node in reversed(stack):
        if node.node_type in anchor_types:
            return max(base_level, node.level + 1)
    return base_level


def _article_context_level(stack: list[StructureNode]) -> int:
    article_level = _current_article_level(stack)
    # 2 matches the level the `numbered_section` branch already applies a
    # staged heading at (`apply_pending_heading(level=2)`) — keeps any real
    # heading directly above the level-3 tier `_article_unit_level` uses
    # below for top-level numbered items with no MADDE article in scope.
    return article_level + 1 if article_level is not None else 2


def _article_unit_level(stack: list[StructureNode], *, kind: str) -> int:
    article_level = _current_article_level(stack)
    if article_level is not None:
        context_levels = [
            node.level
            for node in stack
            if node.level > article_level and node.node_type in {"context", "heading"}
        ]
        if context_levels:
            return max(context_levels) + 1
        return article_level + 1

    # No MADDE article in scope (protocols/agreements numbered "1)", "2)",
    # ... instead of MADDE): top-level paragraph units share the fixed level
    # the `numbered_section` branch above already uses (3), so "1)", "2)",
    # ... are recognized as siblings of each other regardless of which one
    # is currently on the stack — `push_node`'s pop-while-level>=new-level
    # then correctly pops the previous sibling instead of nesting under it.
    # A clause ("a)") nests one level beneath whichever top-level item is
    # currently open.
    return 4 if kind == "clause" else 3


def _current_article_level(stack: list[StructureNode]) -> int | None:
    for node in reversed(stack):
        if node.node_type == "article":
            return node.level
    return None


def _classify_block(block: SourceBlock) -> dict[str, Any]:
    clean = _clean_inline_markdown(block.text)
    folded = _fold_text(clean)

    appendix = _APPENDIX_RE.match(folded)
    if appendix:
        label = appendix.group("num") or clean
        return {"kind": "appendix", "label": f"EK {label}".strip()}

    article = _ARTICLE_RE.match(folded)
    if article:
        label = str(article.group("label"))
        article_no = article.group("num").upper()
        if "gecici" in label:
            article_heading = f"GEÇİCİ MADDE {article_no}"
            article_no = f"GEÇİCİ {article_no}"
        elif "mukerrer" in label:
            article_heading = f"MÜKERRER MADDE {article_no}"
            article_no = f"MÜKERRER {article_no}"
        else:
            article_heading = f"MADDE {article_no}"
        return {
            "kind": "article",
            "article_no": article_no,
            "article_heading": article_heading,
            "rest": _article_rest_from_original(clean),
        }

    if _SECTION_RE.match(folded):
        return {"kind": "section"}

    numbered = _NUMBERED_SECTION_RE.match(clean)
    if numbered and _is_numbered_section_title(clean):
        return {"kind": "numbered_section"}

    if block.kind == "heading" or _is_short_bold_line(block.text):
        return {"kind": "heading"}

    return {"kind": "paragraph"}


def _article_rest_from_original(clean: str) -> str:
    match = re.match(
        r"^(?:GEÇİCİ\s+MADDE|Geçici\s+Madde|MÜKERRER\s+MADDE|Mükerrer\s+Madde|"
        r"MADDE|Madde)\s+\d+[A-Za-z]?\s*(?:[-–—:]\s*)?(.*)$",
        clean,
    )
    return match.group(1).strip() if match else ""


def _article_title_from_rest(rest: str, pending_title: str | None) -> str | None:
    if pending_title and _is_probable_article_title(pending_title):
        return pending_title
    stripped = rest.strip()
    if not stripped:
        return None
    if stripped.startswith("("):
        return pending_title if pending_title else None
    if len(stripped) <= 90 and not stripped.endswith("."):
        return stripped
    return pending_title if pending_title else None


def _updated_section_path(existing: list[str], heading: str) -> list[str]:
    folded = _fold_text(heading)
    if "kisim" in folded:
        return [heading]
    if "bolum" in folded:
        without_old_sections = [
            item
            for item in existing
            if "bolum" not in _fold_text(item) and "kisim" in _fold_text(item)
        ]
        return [*without_old_sections, heading]
    return [*existing, heading]


def _merge_tiny_drafts(drafts: list[_ChunkDraft], min_chars: int) -> list[_ChunkDraft]:
    if not drafts or min_chars <= 0:
        return drafts

    merged: list[_ChunkDraft] = []
    for draft in drafts:
        if (
            merged
            and len(_draft_text(draft)) < min_chars
            and draft.chunk_type in {"free_text", "heading_section"}
            and merged[-1].chunk_type
            in {"free_text", "heading_section", "numbered_section"}
        ):
            previous = merged[-1]
            previous.blocks.extend(draft.blocks)
            previous.warnings.extend(draft.warnings)
            continue
        merged.append(draft)
    return merged


def _drop_non_atomic_parent_drafts(
    drafts: list[_ChunkDraft], metadata: DocumentMetadata
) -> list[_ChunkDraft]:
    if not drafts:
        return drafts

    ancestor_ids_with_chunks: set[str] = set()
    for draft in drafts:
        if _is_non_chunk_draft(draft, metadata) or not _draft_text(draft).strip():
            continue
        ancestor_ids_with_chunks.update(draft.parent_ids[:-1])

    filtered: list[_ChunkDraft] = []
    for draft in drafts:
        if (
            draft.chunk_type == "article"
            and draft.parent_id is not None
            and draft.parent_id in ancestor_ids_with_chunks
        ):
            continue
        filtered.append(draft)
    return filtered


def _is_non_chunk_draft(draft: _ChunkDraft, metadata: DocumentMetadata) -> bool:
    text = _draft_text(draft).strip()
    if not text:
        return True
    if (
        draft.chunk_type == "article"
        and len(draft.blocks) == 1
        and draft.parent_path
        and _same_label(text, draft.parent_path[-1])
    ):
        return True
    if (
        draft.chunk_type == "free_text"
        and len(draft.blocks) == 1
        and metadata.title
        and _same_label(text, metadata.title)
    ):
        return True
    return False


def _article_context_label(draft: _ChunkDraft) -> tuple[str, int | None] | None:
    if draft.chunk_type != "article" or len(draft.blocks) <= 1:
        return None

    context_blocks = draft.blocks[1:]
    text = _clean_inline_markdown(_blocks_text(context_blocks))
    if not text:
        return None
    return _label_tail(text, limit=140), context_blocks[0].block_id


def _preamble_context_label(draft: _ChunkDraft) -> tuple[str, int | None] | None:
    if draft.chunk_type != "free_text" or not draft.blocks:
        return None

    folded_path = _fold_text(" ".join(draft.parent_path))
    folded_text = _fold_text(_draft_text(draft))
    if "akit taraf" not in folded_path or "anlasmislardir" not in folded_text:
        return None

    heading = next(
        (
            item
            for item in reversed(draft.parent_path)
            if "akit taraf" in _fold_text(item)
        ),
        None,
    )
    terminal = next(
        (
            _clean_inline_markdown(block.text)
            for block in reversed(draft.blocks)
            if "anlasmislardir" in _fold_text(block.text)
        ),
        _clean_inline_markdown(draft.blocks[-1].text),
    )
    label = f"{heading} {terminal}" if heading else terminal
    return _label_tail(label, limit=180), draft.blocks[0].block_id


def _paragraph_no(text: str) -> str | None:
    match = _PARAGRAPH_RE.match(text)
    return match.group("num") if match else None


def _clause_label(text: str) -> str | None:
    match = _CLAUSE_RE.match(text)
    return match.group("label") if match else None


def _subclause_label(text: str) -> str | None:
    match = _SUBCLAUSE_RE.match(text)
    return match.group("label") if match else None


def _article_unit(clean_text: str, raw_text: str) -> dict[str, str] | None:
    paragraph = _PARAGRAPH_RE.match(clean_text)
    if paragraph:
        body = clean_text[paragraph.end() :].strip()
        if not body:
            return None
        return {
            "kind": "paragraph",
            "label": paragraph.group("num"),
            "marker": paragraph.group(0).strip(),
            "body": body,
        }

    repeated = _REPEATED_PARAGRAPH_RE.match(clean_text)
    if repeated:
        body = clean_text[repeated.end() :].strip()
        if not body:
            return None
        return {
            "kind": "paragraph",
            "label": f"MÜKERRER {repeated.group('num')}",
            "marker": repeated.group("marker").strip(),
            "body": body,
        }

    clause = _CLAUSE_RE.match(clean_text)
    if not clause:
        return None

    label = clause.group("label")
    marker = clause.group(0).strip()
    is_parenthesized = marker.startswith("(")
    if (
        is_parenthesized
        and _is_roman_label(label)
        and not _starts_with_bold_clause_marker(raw_text, label)
    ):
        return None

    return {
        "kind": "clause",
        "label": _fold_text(label),
        "marker": marker,
        "body": clean_text[clause.end() :].strip(),
    }


def _article_unit_parent_label(unit: dict[str, str], clean_text: str) -> str:
    marker = unit.get("marker") or unit.get("label") or ""
    body = unit.get("body") or clean_text
    label_tail = _label_tail(body)
    return f"{marker} {label_tail}".strip() if label_tail else marker


def _is_structural_article_unit_heading(unit: dict[str, str]) -> bool:
    body = (unit.get("body") or "").strip()
    if not body or len(body) > 90:
        return False
    if '"' in body or "“" in body or "”" in body:
        return False
    if re.search(r"[.;]\s*$", body):
        return False
    return _is_mostly_upper(body)


def _label_tail(text: str, *, limit: int = 90) -> str:
    quote = re.search(r"[\"“](.+?)[\"”]", text)
    if quote:
        tail = f'"{quote.group(1).strip()}"'
    else:
        tail = text.strip()
        tail = re.split(r"[.;]\s+", tail, maxsplit=1)[0].strip()

    tail = re.sub(r"\s+", " ", tail).strip(" -–—:")
    if len(tail) <= limit:
        return tail
    return f"{tail[: limit - 3].rstrip()}..."


def _starts_with_bold_clause_marker(raw_text: str, label: str) -> bool:
    escaped = re.escape(label)
    stripped = raw_text.lstrip()
    return bool(
        re.match(rf"^\*{{1,3}}\s*\({escaped}\)", stripped, flags=re.IGNORECASE)
        or re.match(rf"^<strong>\s*\({escaped}\)", stripped, flags=re.IGNORECASE)
    )


def _is_roman_label(label: str) -> bool:
    return _fold_text(label) in _ROMAN_LABELS


def _is_mostly_upper(text: str) -> bool:
    letters = [char for char in text if char.isalpha()]
    if len(letters) < 4:
        return False
    uppercase = sum(1 for char in letters if char.isupper())
    return uppercase / len(letters) >= 0.75


def _is_numbered_section_title(text: str) -> bool:
    if len(text) > 180:
        return False
    if text.rstrip().endswith(":"):
        return True
    return _is_short_bold_line(text)


def _is_probable_article_title(text: str) -> bool:
    if not text or len(text) > 90:
        return False
    folded = _fold_text(text)
    if _SECTION_RE.match(folded) or _APPENDIX_RE.match(folded):
        return False
    if folded.startswith(("madde ", "gecici madde", "mukerrer madde")):
        return False
    unit = _article_unit(text, text)
    if unit and _is_structural_article_unit_heading(unit):
        return False
    if _NORM_NUMBER_RE.fullmatch(text.strip()):
        return False
    if re.fullmatch(r"\d{1,2}[./]\d{1,2}[./]\d{4}(?:\s*/\s*\d+)?", text.strip()):
        return False
    if _DATE_DMY_RE.search(text):
        return False
    return True


def _looks_like_markdown_heading(text: str) -> bool:
    if text.startswith("#"):
        return True
    return _is_short_bold_line(text)


def _is_short_bold_line(text: str) -> bool:
    stripped = text.strip()
    if len(_clean_inline_markdown(stripped)) > 220:
        return False
    return _is_wrapped_by_marker(stripped, "***") or _is_wrapped_by_marker(
        stripped, "**"
    )


def _is_wrapped_by_marker(text: str, marker: str) -> bool:
    if not text.startswith(marker) or not text.endswith(marker):
        return False
    inner = text[len(marker) : -len(marker)].strip()
    return bool(inner)


def _clean_inline_markdown(text: str) -> str:
    cleaned = text.strip()
    cleaned = re.sub(r"^#+\s*", "", cleaned)
    cleaned = cleaned.replace("\t", " ")
    cleaned = re.sub(r"[*_`]+", "", cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned.strip()


def _fold_text(text: str) -> str:
    normalized = unicodedata.normalize("NFKD", text.casefold())
    without_marks = "".join(ch for ch in normalized if not unicodedata.combining(ch))
    return (
        without_marks.replace("ı", "i")
        .replace("ğ", "g")
        .replace("ü", "u")
        .replace("ş", "s")
        .replace("ö", "o")
        .replace("ç", "c")
    )


def _draft_text(draft: _ChunkDraft) -> str:
    return _blocks_text(draft.blocks)


def _blocks_text(blocks: list[SourceBlock]) -> str:
    return "\n".join(block.text.strip() for block in blocks if block.text.strip())


def _stable_id(prefix: str, value: str) -> str:
    digest = hashlib.sha1(value.encode("utf-8")).hexdigest()
    return f"{prefix}_{digest}"


def chunk_file(
    file_path: str,
    *,
    root_path: str | None = None,
    max_chunk_chars: int = 2400,
) -> ChunkedDocument:
    """Convenience wrapper used by the standalone inspector."""

    return RegulatoryChunker(max_chunk_chars=max_chunk_chars).chunk_file(
        file_path, root_path=root_path
    )
